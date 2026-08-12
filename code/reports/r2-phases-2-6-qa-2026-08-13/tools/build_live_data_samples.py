from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[3]
DOCX_OUT = ROOT / "output" / "r2-live-data" / "devnorth-2026-project-brief.docx"
PDF_OUT = ROOT / "output" / "pdf" / "devnorth-2026-venue-evidence.pdf"


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_docx() -> None:
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(25, 46, 71)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(25, 46, 71)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("DevNorth 2026 — Project Brief")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Approved planning baseline · 13 August 2026")
    run.italic = True
    run.font.color.rgb = RGBColor(85, 105, 125)

    document.add_heading("Outcome", level=1)
    document.add_paragraph(
        "Run DevNorth 2026 as a sold-out, well-rated one-day developer conference "
        "on 18 September 2026 for 450 attendees. Success means at least 405 checked-in "
        "attendees, an average satisfaction score of 4.4/5, and sponsor revenue of PKR 12,000,000."
    )

    document.add_heading("Scope and constraints", level=1)
    for item in (
        "Venue: Expo Centre Lahore, Hall B; contracted capacity 500.",
        "Programme: three tracks, 18 sessions, one keynote, and an evening networking reception.",
        "Budget ceiling: PKR 28,000,000 with a 10% contingency reserve.",
        "Registration closes 11 September 2026; final catering count is due 12 September.",
        "Wi-Fi must support 450 concurrent devices with a tested failover connection.",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Delivery plan", level=1)
    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ("Deliverable", "Owner", "Due", "Acceptance evidence")
    for index, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], value, bold=True)
        table.rows[0].cells[index]._tc.get_or_add_tcPr().append(
            __import__("docx").oxml.parse_xml(
                '<w:shd {} w:fill="193047"/>'.format(__import__("docx").oxml.ns.nsdecls("w"))
            )
        )
        table.rows[0].cells[index].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    rows = (
        ("Venue readiness", "Aisha Khan", "05 Sep", "Capacity and failover test signed by venue manager"),
        ("Registration operations", "Omar Siddiqui", "11 Sep", "405+ confirmed attendees and QR check-in rehearsal"),
        ("Programme readiness", "Sara Ali", "09 Sep", "All speakers confirmed; backup keynote brief approved"),
        ("Sponsor delivery", "Bilal Ahmed", "12 Sep", "PKR 12m contracted and sponsor inventory reconciled"),
    )
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], value)

    document.add_heading("Outcome checkpoints", level=1)
    document.add_paragraph(
        "Checkpoint 1 — 28 August: confirm venue capacity and tested network failover. "
        "Checkpoint 2 — 11 September: confirm 405 registrations, all speaker backups, and sponsor revenue. "
        "If either checkpoint fails, the executive sponsor will reduce programme scope or move to the backup venue."
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("DevNorth 2026 · controlled QA sample")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)
    document.save(DOCX_OUT)


def build_pdf() -> None:
    PDF_OUT.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="QAHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#193047"), alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="QABody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=15, textColor=colors.HexColor("#27384A")))
    doc = SimpleDocTemplate(str(PDF_OUT), pagesize=A4, rightMargin=20 * mm, leftMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    story = [
        Paragraph("DevNorth 2026 — Venue & Network Evidence", styles["QAHeading"]),
        Paragraph("Issued by Expo Centre Lahore · Hall B Operations · evidence reference VNE-2026-0918", styles["QABody"]),
        Spacer(1, 8),
    ]
    data = [
        ["Evidence field", "Verified value"],
        ["Venue", "Expo Centre Lahore — Hall B"],
        ["Safe seated capacity", "500 attendees"],
        ["Primary network", "2.0 Gbps dedicated fibre"],
        ["Concurrent device test", "500 devices sustained for 45 minutes"],
        ["Failover", "500 Mbps secondary fibre; automatic cutover under 60 seconds"],
        ["Test date", "28 August 2026"],
        ["Accountable owner", "Aisha Khan, Venue Operations Lead"],
        ["Acceptance", "Passed — supports the 450-device DevNorth requirement"],
    ]
    table = Table(data, colWidths=[55 * mm, 95 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#193047")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEADING", (0, 0), (-1, -1), 13),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C4CF")),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F5F7F9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([table, Spacer(1, 12), Paragraph("Control note: this document is a realistic QA fixture generated for OSLO Release 2 extraction and reanalysis verification. It contains no customer or personal production data.", styles["QABody"])])
    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_OUT)
    print(PDF_OUT)
